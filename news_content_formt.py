import os
import re
import tkinter as tk
import traceback
from tkinter import ttk, messagebox
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.shared import Pt, Cm, Mm
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from datetime import datetime


class ArticleGeneratorApp:
    def __init__(self, master):
        self.master = master
        master.title("文章生成器")
        master.geometry("500x400")

        # 创建主框架
        main_frame = ttk.Frame(master, padding="10")
        main_frame.grid(row=0, column=0, sticky=(tk.W, tk.E, tk.N, tk.S))

        # 创建并放置输入框和标签
        self.create_input_field(main_frame, 0, "标题:", "title")
        self.create_input_field(main_frame, 1, "栏目:", "category")
        self.create_input_field(main_frame, 2, "来源:", "source")
        self.create_input_field(main_frame, 3, "发布日期:", "publish_date")

        # 创建正文输入框
        ttk.Label(main_frame, text="正文:").grid(row=4, column=0, sticky=tk.W, pady=5)
        self.content_text = tk.Text(main_frame, height=10, width=50)
        self.content_text.grid(row=5, column=0, columnspan=2, pady=5)

        # 创建生成按钮
        generate_button = ttk.Button(main_frame, text="生成docx文档", command=self.generate_docx)
        generate_button.grid(row=6, column=0, columnspan=2, pady=10)

        # 配置列和行的权重，使其可以随窗口调整大小
        master.columnconfigure(0, weight=1)
        master.rowconfigure(0, weight=1)
        main_frame.columnconfigure(1, weight=1)

    def create_input_field(self, parent, row, label_text, attr_name):
        ttk.Label(parent, text=label_text).grid(row=row, column=0, sticky=tk.W, pady=5)
        entry = ttk.Entry(parent)
        entry.grid(row=row, column=1, sticky=(tk.W, tk.E), pady=5)
        setattr(self, f"{attr_name}_entry", entry)

    def add_formatted_paragraph(self, doc, text, style, center=False):
        paragraph = doc.add_paragraph(style=style)
        if center:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # 更新正则表达式以匹配数字、英文字母、百分号和小数点
        parts = re.split(r'(\d+\.?\d*%?|\.|%|[a-zA-Z]+)', text)

        for part in parts:
            if part.strip():
                if '%' in part:
                    # 处理包含百分号的部分
                    sub_parts = part.split('%')
                    for i, sub_part in enumerate(sub_parts):
                        if sub_part:
                            run = paragraph.add_run(sub_part)
                            run.style = 'Custom Number Letter'
                        if i < len(sub_parts) - 1:  # 不是最后一个部分，添加百分号
                            run = paragraph.add_run('%')
                            run.font.name = '仿宋_GB2312'
                            run._element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')
                elif re.match(r'\d+\.?\d*', part) or re.match(r'[a-zA-Z]+', part) or part == '.':
                    # 数字、英文字母或单独的小数点
                    run = paragraph.add_run(part)
                    run.style = 'Custom Number Letter'
                else:
                    # 其他文本
                    run = paragraph.add_run(part)
                    # 确保其他文本使用段落的默认样式
                    run.style = None

        # 设置段落的东亚字体
        for run in paragraph.runs:
            if run.style is None:
                run.font.name = paragraph.style.font.name
                run._element.rPr.rFonts.set(qn('w:eastAsia'), paragraph.style.font.name)

    def generate_docx(self):
        # 获取输入的值
        title = self.title_entry.get()
        category = self.category_entry.get()
        source = self.source_entry.get()
        publish_date = self.publish_date_entry.get()
        content = self.content_text.get("1.0", tk.END).strip()

        # 验证输入
        if not all([title, category, source, publish_date, content]):
            messagebox.showerror("错误", "请填写所有字段")
            return

        # 创建文档
        doc = Document()
        # 设置纸张大小为A4
        section = doc.sections[0]
        section.page_height = Mm(297)
        section.page_width = Mm(210)
        # 设置页边距
        sections = doc.sections
        for section in sections:
            section.top_margin = Cm(3.7)
            section.bottom_margin = Cm(3.5)
            section.left_margin = Cm(2.8)
            section.right_margin = Cm(2.6)

        # 定义字体样式
        self.define_styles(doc)

        # 添加页码
        self.add_page_number(doc)
        # 添加标题
        self.add_formatted_paragraph(doc, title, 'Custom Title', center=True)

        # 添加栏目、来源、发布日期
        self.add_formatted_paragraph(doc, f"栏目: {category}", 'Custom Subtitle')
        self.add_formatted_paragraph(doc, f"来源: {source}", 'Custom Subtitle')
        self.add_formatted_paragraph(doc, f"发布日期: {publish_date}", 'Custom Subtitle')

        # 添加空行
        doc.add_paragraph(style='Custom Subtitle')
        # 添加正文
        paragraphs = content.split('\n')
        for paragraph in paragraphs:
            if paragraph.strip():  # 忽略空行
                self.add_formatted_paragraph(doc, paragraph.strip(), 'Custom Content')

        # 保存文档
        try:
            filename = f"{title}_{datetime.now().strftime('%Y%m%d%H%M%S')}.docx"
            # 创建文件夹，名字为当前日期
            today = str(datetime.now().date())
            os.makedirs(today, exist_ok=True)
            file_path = os.path.join(today, filename)
            doc.save(file_path)
            messagebox.showinfo("成功", f"文档已保存为 {filename}")
        except Exception as e:
            print(traceback.format_exc())
            messagebox.showerror("错误", f"保存文档时出错: {str(e)}")

        # 清空缓存
        self.title_entry.delete(0, tk.END)
        self.category_entry.delete(0, tk.END)
        self.source_entry.delete(0, tk.END)
        self.publish_date_entry.delete(0, tk.END)
        self.content_text.delete("1.0", tk.END)

    def define_styles(self, doc):
        # 通用段落格式设置函数
        def set_paragraph_format(style):
            style.paragraph_format.line_spacing = Pt(28)
            style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
            style.paragraph_format.space_after = Pt(0)  # 设置段后间距为0磅

        # 标题样式
        if 'Custom Title' not in doc.styles:
            title_style = doc.styles.add_style('Custom Title', WD_STYLE_TYPE.PARAGRAPH)
        else:
            title_style = doc.styles['Custom Title']
        title_style.font.name = '方正小标宋简体'
        title_style.font.size = Pt(22)  # 二号字体大小约为29磅
        title_style._element.rPr.rFonts.set(qn('w:eastAsia'), '方正小标宋简体')
        set_paragraph_format(title_style)

        # 副标题样式（用于栏目、来源、发布日期）
        if 'Custom Subtitle' not in doc.styles:
            subtitle_style = doc.styles.add_style('Custom Subtitle', WD_STYLE_TYPE.PARAGRAPH)
        else:
            subtitle_style = doc.styles['Custom Subtitle']
        subtitle_style.font.name = '黑体'
        subtitle_style.font.size = Pt(16)  # 三号字体大小约为16磅
        subtitle_style._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
        set_paragraph_format(subtitle_style)

        # 正文样式
        if 'Custom Content' not in doc.styles:
            content_style = doc.styles.add_style('Custom Content', WD_STYLE_TYPE.PARAGRAPH)
        else:
            content_style = doc.styles['Custom Content']
        content_style.font.name = '仿宋_GB2312'
        content_style.font.size = Pt(16)  # 三号字体大小约为16磅
        content_style._element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')
        content_style.paragraph_format.first_line_indent = 0
        content_style.paragraph_format.element.pPr.ind.set(qn("w:firstLineChars"), '200')
        set_paragraph_format(content_style)

        # 数字和英文字母样式
        if 'Custom Number Letter' not in doc.styles:
            number_letter_style = doc.styles.add_style('Custom Number Letter', WD_STYLE_TYPE.CHARACTER)
        else:
            number_letter_style = doc.styles['Custom Number Letter']
        number_letter_style.font.name = 'Times New Roman'

    def add_page_number(self, doc):
        sections = doc.sections
        for section in sections:
            footer = section.footer
            footer_para = footer.paragraphs[0]
            footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

            run = footer_para.add_run()
            fldChar = OxmlElement('w:fldChar')
            fldChar.set(qn('w:fldCharType'), 'begin')

            instrText = OxmlElement('w:instrText')
            instrText.set(qn('xml:space'), 'preserve')
            instrText.text = 'PAGE'

            fldChar2 = OxmlElement('w:fldChar')
            fldChar2.set(qn('w:fldCharType'), 'end')

            run._element.append(fldChar)
            run._element.append(instrText)
            run._element.append(fldChar2)

            run.font.name = 'Times New Roman'
            run.font.size = Pt(9)  # 小五号字体大小约为9磅


def main():
    root = tk.Tk()
    app = ArticleGeneratorApp(root)
    root.mainloop()


if __name__ == "__main__":
    main()
